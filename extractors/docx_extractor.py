# extractor/docx_extractor.py
"""
DOCX/DOC Text Extractor
Extracts text from Microsoft Word documents (.docx, .doc)
"""

from docx import Document
import requests
from io import BytesIO
import os


def extract_text_from_docx_file(file_path):
    """
    Extract text from local DOCX file
    
    Args:
        file_path (str): Full path to DOCX file
        
    Returns:
        tuple: (success: bool, text: str, error: str)
    """
    try:
        if not os.path.exists(file_path):
            return False, "", f"File not found: {file_path}"
        
        
        doc = Document(file_path)
        
        text = ""
        
        # Extract text from paragraphs
        
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        if len(text.strip()) < 10:
            return False, "", "DOCX appears to be empty or contains no text"
        
        return True, text.strip(), None
        
    except Exception as e:
        return False, "", f"Error extracting DOCX: {str(e)}"


def extract_text_from_docx_url(url):
    """
    Extract text from DOCX file via URL
    
    Args:
        url (str): URL to DOCX file
        
    Returns:
        tuple: (success: bool, text: str, error: str)
    """
    try:
        response = requests.get(url, timeout=30)
        response.raise_for_status()
        
        # Open DOCX from bytes
        docx_file = BytesIO(response.content)
        doc = Document(docx_file)
        
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        if len(text.strip()) < 10:
            return False, "", "DOCX appears to be empty or contains no text"
        
        return True, text.strip(), None
        
    except requests.exceptions.RequestException as e:
        return False, "", f"Error downloading DOCX: {str(e)}"
    except Exception as e:
        return False, "", f"Error extracting DOCX from URL: {str(e)}"


def extract_text_from_docx(docx_path, base_path=""):
    """
    Extract text from DOCX - handles both local files and URLs
    
    Args:
        docx_path (str): DOCX filename or relative path
        base_path (str): Base path (folder path or URL)
        
    Returns:
        tuple: (success: bool, text: str, error: str)
    """
    # Build full path
    if base_path:
        if base_path.endswith('/') or base_path.endswith('\\'):
            full_path = base_path + docx_path
        else:
            full_path = os.path.join(base_path, docx_path)
    else:
        full_path = docx_path
    
    # print(f"[DOCX] Extracting DOCX: {os.path.basename(docx_path)}")
    
    # Check if it's a URL or local file
    if full_path.startswith('http://') or full_path.startswith('https://'):
        return extract_text_from_docx_url(full_path)
    else:
        return extract_text_from_docx_file(full_path)


def extract_text_from_doc(doc_path, base_path=""):
    """
    Extract text from old DOC format (requires conversion or different library)
    Note: python-docx doesn't support .doc format directly
    
    Args:
        doc_path (str): DOC filename
        base_path (str): Base path
        
    Returns:
        tuple: (success: bool, text: str, error: str)
    """
    return False, "", "Old .doc format not supported. Please convert to .docx or use a different extractor."